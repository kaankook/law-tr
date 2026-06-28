#python src/data_ingestion.py
from __future__ import annotations
import asyncio
import hashlib
import json
import logging
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

_PROJECT_ROOT = Path(__file__).parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

from docx import Document
from langchain_core.documents import Document as LangchainDocument
from openai import AsyncOpenAI
try:
    from src.config import (
        CHUNK_CFG,
        JSON_FILE_PATH,
        LLM_ENRICHMENT_CONCURRENCY,
        LLM_ENRICHMENT_MODEL,
        LLM_ENRICHMENT_PROMPT,
        LLM_ENRICHMENT_RETRIES,
        MAPPING_FILE_PATH,
        PROCESSED_DIR,
        RAW_DATA_DIR,
        SOURCE_FORMAT_MAPPING,
    )
except ImportError as e:
    raise ImportError(f"Configuration import failed: {e}") from e

from src.law_mapping_resolver import LawMappingResolver

logger = logging.getLogger(__name__)


class MetadataKey:
    # Source information
    SOURCE            = "source"
    DOSYA_ADI         = "dosya_adi"
    LAW_NUMBER        = "law_number"

    # Document structure
    ARTICLE_REFERENCE = "article_reference"
    ARTICLE_NUMBER    = "article_number"
    ARTICLE_TYPE      = "article_type"

    # Chunk information
    CHUNK_ID          = "chunk_id"
    CHUNK_INDEX       = "chunk_index"
    CHUNK_TOTAL       = "chunk_total"
    CHUNK_PART        = "chunk_part"
    PREV_CHUNK_ID     = "prev_chunk_id"
    NEXT_CHUNK_ID     = "next_chunk_id"

    # Mülga (deprecated) information
    IS_MULGA_SOURCE   = "is_mulga_source"
    IS_PARTIAL_MULGA  = "is_partial_mulga"
    MAPS_TO_LAW_NO    = "maps_to_law_no"
    MAPS_TO_LAW_NAME  = "maps_to_law_name"
    MULGA_SCOPE       = "mulga_scope"
    MULGA_EXCEPTION   = "mulga_exception"

    # Size metrics
    CHAR_COUNT        = "char_count"
    WORD_COUNT        = "word_count"

# ─────────────────────────────────────────────
# REGEX PAREKSİLERİ
# ─────────────────────────────────────────────
_RE_ARTICLE = re.compile(
    r"^\s*"
    r"(?P<type>Geçici\s+Madde|GEÇİCİ\s+MADDE|Ek\s+Madde|EK\s+MADDE|MADDE|Madde)\s*"
    r"(?P<no>\d+[A-Za-z0-9/]*)"
    r"(?:\s*[-–—:/]\s*(?P<rest>.*))?$",
    re.IGNORECASE | re.DOTALL,
)
_RE_ARTICLE_BARE = re.compile(
    r"^\s*(?P<type>Madde|MADDE)\s*(?P<no>\d+[A-Za-z0-9/]*)\s*$",
    re.IGNORECASE,
)
_RE_FIKRA        = re.compile(r"^\s*\((\d+)\)\s*(.*)$", re.DOTALL)
_RE_BENT         = re.compile(r"^\s*[a-zA-ZçşğüöıÇŞĞÜÖİ]\)\s", re.UNICODE)
_RE_ROMAN        = re.compile(
    r"^\s*(I{1,3}|IV|VI{0,3}|IX|X{1,3})\s*[\.\-–]\s*\S",
    re.IGNORECASE,
)

_RE_FOOTNOTE_LINE   = re.compile(r"^\s*\[\^[\w-]+\]:?\s*$", re.MULTILINE)
_RE_INLINE_FOOTNOTE = re.compile(r"\^\[?\d+\]?\^")
_RE_CANCELLED       = re.compile(r"\((?:\.{2,3}|…)\)")
_CANCELLED_TOKEN    = "__MULGA_CANCELLED__"

_RE_META = re.compile(
    r"^(Kanun\s+Numara|Kabul\s+Tarihi|Yayımlandığı|Resmî\s+Gazete|Düstur)",
    re.IGNORECASE,
)

_RE_GECICI_MADDE = re.compile(r"\bgecici\s+madde\b", re.IGNORECASE)
_RE_EK_MADDE     = re.compile(r"\bek\s+madde\b", re.IGNORECASE)


# ─────────────────────────────────────────────
# YARDIMCI FONKSİYONLAR
# ─────────────────────────────────────────────
def _clean(text: str) -> str:
    """Metni normalize et: boşlukları temizle, bölünmüş kelimeler birleştir."""
    text = re.sub(r"[\u200b\u200c\u200d\ufeff\xa0]", " ", text)
    text = _RE_FOOTNOTE_LINE.sub("", text)
    text = _RE_INLINE_FOOTNOTE.sub("", text)
    text = _RE_CANCELLED.sub(f" {_CANCELLED_TOKEN} ", text)
    text = re.sub(r"([a-zışğüöçıA-ZİĞÜŞÖÇ])-\s+([a-zA-ZışğüöçıİĞÜŞÖÇ])", r"\1\2", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def _md5(s: str, n: int = 16) -> str:
    return hashlib.md5(s.encode()).hexdigest()[:n]

def _article_ref(raw_type: str, no: str) -> str:
    """Madde türüne göre referans formatı döndür (Geçici/Ek/Normal)."""
    t        = raw_type.lower()
    no_clean = str(no).strip() if no else ""
    if _RE_GECICI_MADDE.search(t):
        return f"Geçici Madde {no_clean}".rstrip()
    if _RE_EK_MADDE.search(t):
        return f"Ek Madde {no_clean}".rstrip()
    return f"Madde {no_clean}".rstrip()

def _article_type(raw_type: str) -> str:
    """Madde türünü belirle (gecici_madde/ek_madde/madde)."""
    t = raw_type.lower()
    if _RE_GECICI_MADDE.search(t):
        return "gecici_madde"
    if _RE_EK_MADDE.search(t):
        return "ek_madde"
    return "madde"

def _chunk_id(source: str, ref: str, idx: int) -> str:
    """Chunk için benzersiz ID üret: safe_name_index_hash."""
    h    = _md5(f"{source}||{ref}||{idx}", 16)
    safe = re.sub(r"[\s\W]", "_", f"{source}_{ref}", flags=re.UNICODE)[:50]
    return f"{safe}_{idx}_{h}"

def _is_empty_fikra(text: str) -> bool:
    """Fıkra satırı boş mu kontrol et."""
    m = _RE_FIKRA.match(text)
    return bool(m) and (not m.group(2) or not m.group(2).strip())


# ─────────────────────────────────────────────
# BLOK YAPISI
# ─────────────────────────────────────────────
@dataclass
class _Block:
    """Metinsellik bloğu: BENT (a), b), ...), ROMAN (I., II.), FIKRA (1), (2) grupları."""
    lines: List[str] = field(default_factory=list)
    kind: str = "fikra"

    def add(self, line: str) -> None:
        self.lines.append(line)

    @property
    def text(self) -> str:
        return "\n".join(self.lines)

    def charlen(self) -> int:
        return sum(len(l) for l in self.lines) + max(0, len(self.lines) - 1)

    def empty(self) -> bool:
        return not self.lines

def _is_closing_clause(line: str, prev: str) -> bool:
    """Satır, BENT maddesi devamı mı kontrol et."""
    if not _RE_BENT.match(prev):
        return False
    s = line.strip()
    if not s:
        return False
    return s[0].islower() or len(s) < 80

def _build_blocks(lines: List[str]) -> List[_Block]:
    """Satırları yapısal bloklar halinde grupla: BENT, ROMAN, FIKRA veya metin."""
    blocks: List[_Block] = []
    cur = _Block(kind="fikra")

    for raw in lines:
        s = raw.strip()
        if not s:
            continue

        if _RE_BENT.match(s):
            if cur.kind != "bent_group" and not cur.empty():
                blocks.append(cur)
                cur = _Block(kind="bent_group")
            elif cur.kind != "bent_group":
                cur.kind = "bent_group"
            cur.add(s)
        elif _RE_ROMAN.match(s):
            if not cur.empty():
                blocks.append(cur)
            cur = _Block(kind="roman")
            cur.add(s)
        elif _RE_FIKRA.match(s):
            if _is_empty_fikra(s):
                continue
            if not cur.empty():
                blocks.append(cur)
            cur = _Block(kind="fikra")
            cur.add(s)
        else:
            prev = cur.lines[-1] if cur.lines else ""
            if prev and _is_closing_clause(s, prev):
                cur.add(s)
            else:
                if not cur.empty():
                    blocks.append(cur)
                cur = _Block(kind="fikra")
                cur.add(s)

    if not cur.empty():
        blocks.append(cur)

    return blocks

def _split_sentences(text: str, target: int) -> List[str]:
    """Metni cümle sınırlarında böl (Türkçe duyarlı)."""
    sents = re.split(r"(?<=[\.\?\!])\s+(?=[A-ZİÇŞĞÜÖ\(])", text)
    chunks, buf = [], ""
    for s in sents:
        if not buf:
            buf = s
        elif len(buf) + 1 + len(s) <= target:
            buf += " " + s
        else:
            chunks.append(buf.strip())
            buf = s
    if buf:
        chunks.append(buf.strip())
    return chunks or [text]

def split_article(
    article_text: str,
    article_ref: str,
    target_size: int,
) -> List[str]:
    """Uzun maddeyi hedef boyuta göre böl, çoklu blok için devam başlığı ekle."""
    text: str = article_text.strip()
    if not text:
        return []
    if len(text) <= target_size:
        return [text]

    blocks: List[_Block] = _build_blocks([l.strip() for l in text.split("\n") if l.strip()])
    if not blocks:
        sentences = _split_sentences(text, target_size)
        return sentences if sentences else [text]

    groups: List[List[_Block]] = []
    cur_grp: List[_Block]      = []
    cur_len: int               = len(article_ref) + 2

    for blk in blocks:
        blen: int = blk.charlen()
        if blen > target_size:
            if cur_grp:
                groups.append(cur_grp)
                cur_grp, cur_len = [], len(article_ref) + 2
            for piece in _split_sentences(blk.text, target_size):
                pb = _Block()
                pb.add(piece)
                groups.append([pb])
        elif cur_len + blen + 2 > target_size and cur_grp:
            groups.append(cur_grp)
            cur_grp = [blk]
            cur_len = len(article_ref) + 2 + blen
        else:
            cur_grp.append(blk)
            cur_len += blen + 2

    if cur_grp:
        groups.append(cur_grp)

    total: int     = len(groups)
    result: List[str] = []
    for i, grp in enumerate(groups):
        body: str = "\n\n".join(b.text for b in grp)
        if i == 0:
            result.append(body)
        else:
            result.append(f"{article_ref} [devam {i+1}/{total}]\n\n{body}")

    return result if result else [text]

# ─────────────────────────────────────────────
# CHUNK OLUŞTURMA
# ─────────────────────────────────────────────
def _make_chunk(
    content: str,
    meta: Dict[str, Any],
    chunk_idx: int,
    chunk_total: int,
    prev_id: Optional[str],
    next_id: Optional[str],
) -> LangchainDocument:
    """İçerik ve metadata'dan LangchainDocument chunk oluştur."""
    meta_safe = dict(meta)
    ctx_parts = [
        f"KAYNAK: {meta_safe['source']}",
        f"MADDE: {meta_safe['article_reference']}",
    ]
    if chunk_total > 1:
        ctx_parts.append(f"PARÇA: {chunk_idx}/{chunk_total}")
    if meta_safe.get(MetadataKey.IS_MULGA_SOURCE):
        ctx_parts.append(f"[MÜLGA → {meta_safe.get(MetadataKey.MAPS_TO_LAW_NAME, '')}]")

    final = (
        f"BAĞLAM: {' | '.join(ctx_parts)}\n\n"
        f"İÇERİK:\n{content}"
    )
    chunk_id = _chunk_id(meta_safe["source"], meta_safe["article_reference"], chunk_idx)

    output_meta: Dict[str, Any] = {
        MetadataKey.SOURCE:            meta_safe[MetadataKey.SOURCE],
        MetadataKey.DOSYA_ADI:         meta_safe[MetadataKey.DOSYA_ADI],
        MetadataKey.LAW_NUMBER:        meta_safe.get(MetadataKey.LAW_NUMBER, ""),
        MetadataKey.ARTICLE_REFERENCE: meta_safe[MetadataKey.ARTICLE_REFERENCE],
        MetadataKey.ARTICLE_NUMBER:    meta_safe.get(MetadataKey.ARTICLE_NUMBER, ""),
        MetadataKey.ARTICLE_TYPE:      meta_safe.get(MetadataKey.ARTICLE_TYPE, "madde"),
        MetadataKey.CHUNK_ID:          chunk_id,
        MetadataKey.CHUNK_INDEX:       chunk_idx,
        MetadataKey.CHUNK_TOTAL:       chunk_total,
        MetadataKey.CHUNK_PART:        f"{chunk_idx}/{chunk_total}",
        MetadataKey.PREV_CHUNK_ID:     prev_id or "",
        MetadataKey.NEXT_CHUNK_ID:     next_id or "",
        MetadataKey.CHAR_COUNT:        len(content),
        MetadataKey.WORD_COUNT:        len(content.split()),
    }

    # Dinamik mülga alanları — sadece mülga kaynak ise eklenir
    if meta_safe.get(MetadataKey.IS_MULGA_SOURCE):
        output_meta[MetadataKey.IS_MULGA_SOURCE]  = True
        output_meta[MetadataKey.MAPS_TO_LAW_NO]   = meta_safe.get(MetadataKey.MAPS_TO_LAW_NO, "")
        output_meta[MetadataKey.MAPS_TO_LAW_NAME] = meta_safe.get(MetadataKey.MAPS_TO_LAW_NAME, "")
        output_meta[MetadataKey.MULGA_SCOPE]       = meta_safe.get(MetadataKey.MULGA_SCOPE, "")
        output_meta[MetadataKey.MULGA_EXCEPTION]   = meta_safe.get(MetadataKey.MULGA_EXCEPTION, "")
    if meta_safe.get(MetadataKey.IS_PARTIAL_MULGA):
        output_meta[MetadataKey.IS_PARTIAL_MULGA]  = True

    return LangchainDocument(page_content=final, metadata=output_meta)

# ─────────────────────────────────────────────
# DOCX NORMALIZE
# ─────────────────────────────────────────────
@dataclass
class _Para:
    """DOCX'ten normalize edilmiş paragraf."""
    style: str
    text:  str

def _normalize(doc: Document) -> List[_Para]:
    """DOCX paragraflarını normalize et, devam satırlarını birleştir."""
    raw: List[_Para] = []
    for i, p in enumerate(doc.paragraphs):
        try:
            t = _clean(p.text)
        except TimeoutError:
            logger.warning(f"Paragraf {i} okunurken zaman aşımı, atlanıyor")
            continue
        except Exception as e:
            logger.warning(f"Paragraf {i} hatasında ({type(e).__name__}), atlanıyor: {e}")
            continue

        if not t or _RE_META.match(t):
            continue

        sty = p.style.name

        if (
            raw
            and raw[-1].style == "List Paragraph"
            and sty in ("Body Text", "Normal")
            and not _RE_ARTICLE.match(t)
            and not _RE_FIKRA.match(t)
            and not _RE_BENT.match(t)
            and len(t) < 80
            and t[0].islower()
        ):
            raw[-1] = _Para(style=raw[-1].style, text=raw[-1].text + " " + t)
            continue

        raw.append(_Para(style=sty, text=t))

    return raw

# ─────────────────────────────────────────────
# MAKALE BUFFER
# ─────────────────────────────────────────────
class _ArticleBuffer:
    """Madde metni biriktir, hazır olunca chunk olarak döndür."""
    def __init__(self, resolver: LawMappingResolver) -> None:
        self._res:  LawMappingResolver = resolver
        self.lines: List[str]          = []
        self.meta:  Dict[str, Any]     = {}

    def reset(self, meta: Dict[str, Any]) -> None:
        self.lines = []
        self.meta  = dict(meta)

    def empty(self) -> bool:
        return not self.lines

    def add(self, text: str) -> None:
        if text:
            self.lines.append(text)

    def flush(self, out: List[LangchainDocument]) -> None:
        if not self.lines:
            return

        full                 = "\n".join(self.lines).strip()
        has_cancelled_marker = _CANCELLED_TOKEN in full
        if has_cancelled_marker:
            full = full.replace(_CANCELLED_TOKEN, "mülga")

        if self.meta.get("article_type") == "giris" and len(full) < 100:
            self.lines = []
            return

        if len(full) < CHUNK_CFG.min_chunk_chars:
            logger.debug(
                "Atlandı (<%d char): %s",
                CHUNK_CFG.min_chunk_chars,
                self.meta.get("article_reference"),
            )
            self.lines = []
            return

        ref    = self.meta.get("article_reference", "")
        law_no = self.meta.get("law_number", "")
        art_no = self.meta.get("article_number", "")

        chunk_meta = dict(self.meta)
        chunk_meta.update(self._res.get_metadata_flags(law_no, art_no))
        if has_cancelled_marker:
            chunk_meta[MetadataKey.IS_PARTIAL_MULGA] = True
            chunk_meta[MetadataKey.IS_MULGA_SOURCE]  = True

        subs  = split_article(full, ref, CHUNK_CFG.target_size) if len(full) > CHUNK_CFG.target_size else [full]
        total = len(subs)

        for i, sub in enumerate(subs):
            idx     = i + 1
            prev_id = _chunk_id(chunk_meta["source"], ref, idx - 1) if i > 0           else None
            next_id = _chunk_id(chunk_meta["source"], ref, idx + 1) if i < total - 1   else None
            out.append(_make_chunk(sub, chunk_meta, idx, total, prev_id, next_id))

        self.lines = []

# ─────────────────────────────────────────────
# YÜKLEME FONKSİYONLARI
# ─────────────────────────────────────────────
def load_and_chunk_legislation(
    file_path: str,
    resolver:  LawMappingResolver,
) -> List[LangchainDocument]:
    """Kanun dosyasını yükle ve maddelere göre böl."""
    if not os.path.exists(file_path):
        logger.info("Dosya bulunamadı (işlenmiyor): %s", file_path)
        return []

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.error("DOCX okunumunda hata: %s — %s", file_path, exc, exc_info=True)
        return []

    fname  = os.path.basename(file_path)
    src    = SOURCE_FORMAT_MAPPING.get(fname)
    name   = src.display_name if src else fname
    law_no = src.law_number   if src else ""

    logger.info("İşleniyor: %s", name)

    try:
        paras = _normalize(doc)
    except Exception as exc:
        logger.error("Paragraflara normalize etme hatasi (%s): %s", fname, exc, exc_info=True)
        return []

    chunks: List[LangchainDocument] = []

    base_meta = {
        "source":            name,
        "dosya_adi":         fname,
        "law_number":        law_no,
        "article_reference": "Giriş",
        "article_number":    "0",
        "article_type":      "giris",
    }

    buf = _ArticleBuffer(resolver)
    buf.reset(base_meta)

    for para in paras:
        t     = para.text
        art_m = _RE_ARTICLE.match(t) or _RE_ARTICLE_BARE.match(t)

        if art_m:
            buf.flush(chunks)

            raw_type = art_m.group("type")
            art_no   = art_m.group("no")
            ref      = _article_ref(raw_type, art_no)

            new_meta = dict(buf.meta or base_meta)
            new_meta.update({
                "article_reference": ref,
                "article_number":    str(art_no),
                "article_type":      _article_type(raw_type),
                "law_number":        law_no,
            })
            buf.reset(new_meta)
            buf.add(t)
            continue

        buf.add(t)

    buf.flush(chunks)
    logger.info("  → %d chunk", len(chunks))
    return chunks

def load_all_documents(
    directory_path: str,
    resolver:       LawMappingResolver,
) -> List[LangchainDocument]:
    """Dizindeki tüm kanun dosyalarını yükle ve böl."""
    dp = Path(directory_path)
    if not dp.exists():
        logger.error("Klasör yok: %s", dp)
        return []

    files = sorted(
        f.name for f in dp.glob("*.docx")
        if not f.name.startswith("~$")
    )
    if not files:
        logger.info("Klasörde .docx dosyası bulunamadı: %s", dp)
        return []

    logger.info("=" * 60)
    logger.info("BELGE İŞLEME — %d dosya", len(files))

    all_docs: List[LangchainDocument] = []
    for fname in files:
        docs = load_and_chunk_legislation(str(dp / fname), resolver)
        all_docs.extend(docs)

    logger.info("=" * 60)
    logger.info("TOPLAM: %d chunk işlendi", len(all_docs))
    return all_docs

# ─────────────────────────────────────────────
# LLM ZENGİNLEŞTİRME
# ─────────────────────────────────────────────
_openai_client = AsyncOpenAI()

def _extract_article_text(page_content: str) -> str:
    """'İÇERİK:' bloğundan sonraki metni çıkar."""
    marker = "İÇERİK:\n"
    pos    = page_content.find(marker)
    return page_content[pos + len(marker):].strip() if pos != -1 else page_content.strip()

def _inject_enrichment(page_content: str, summary: str, questions: List[str]) -> str:
    """Özet ve soruları content'e ekle: ÖZET - SORULAR - İÇERİK sırası."""
    parts: List[str] = []
    if summary:
        parts.append(f"ÖZET: {summary}")
    if questions:
        parts.append("SORULAR:\n" + "\n".join(f"- {q}" for q in questions))
    if not parts:
        return page_content

    block  = "\n\n".join(parts)
    marker = "İÇERİK:\n"
    pos    = page_content.find(marker)
    if pos != -1:
        return f"{page_content[:pos].rstrip()}\n\n{block}\n\n{page_content[pos:]}"
    return f"{block}\n\n{page_content}"

async def _enrich_single(
    doc:     LangchainDocument,
    sem:     asyncio.Semaphore,
    retries: int,
) -> LangchainDocument:
    """Bir chunk'ı LLM'den özet + sorularla zenginleştir (üstel backoff ile retry)."""
    article_ref  = doc.metadata.get("article_reference", "?")
    article_text = _extract_article_text(doc.page_content)

    if len(article_text) < 80:
        logger.debug("Zenginleştirme atlandı (çok kısa, %d char): %s", len(article_text), article_ref)
        return doc

    last_exc: Optional[Exception] = None

    for attempt in range(retries):
        try:
            async with sem:
                response = await _openai_client.chat.completions.create(
                    model=LLM_ENRICHMENT_MODEL,
                    messages=[
                        {"role": "system", "content": LLM_ENRICHMENT_PROMPT},
                        {"role": "user",   "content": article_text},
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.2,
                    max_tokens=512,
                )

            parsed    = json.loads(response.choices[0].message.content or "{}")
            summary:   str        = parsed.get("ozet", "").strip()
            questions: List[str]  = [q.strip() for q in parsed.get("sorular", []) if q.strip()]

            if not summary and not questions:
                logger.warning("Boş LLM yanıtı: %s", article_ref)
                return doc

            new_meta                  = dict(doc.metadata)
            new_meta["llm_summary"]   = summary
            new_meta["llm_questions"] = " | ".join(questions)
            new_meta["llm_enriched"]  = True

            return LangchainDocument(
                page_content=_inject_enrichment(doc.page_content, summary, questions),
                metadata=new_meta,
            )

        except json.JSONDecodeError as exc:
            last_exc = exc
            logger.warning("JSON parse hatası (deneme %d/%d) [%s]: %s", attempt + 1, retries, article_ref, exc)
        except Exception as exc:
            last_exc = exc
            logger.warning("API hatası (deneme %d/%d) [%s]: %s", attempt + 1, retries, article_ref, exc)
            if attempt < retries - 1:
                await asyncio.sleep(2 ** attempt)  # 1s → 2s → 4s

    logger.error("Zenginleştirme başarısız, orijinal chunk korunuyor [%s]: %s", article_ref, last_exc)
    return doc

async def enrich_chunks(
    docs:        List[LangchainDocument],
    concurrency: Optional[int] = None,
    retries:     Optional[int] = None,
) -> List[LangchainDocument]:
    """Tüm chunk'ları LLM ile zenginleştir (özet + sorular)."""
    _concurrency = concurrency if concurrency is not None else LLM_ENRICHMENT_CONCURRENCY
    _retries     = retries     if retries     is not None else LLM_ENRICHMENT_RETRIES

    sem = asyncio.Semaphore(_concurrency)
    logger.info(
        "LLM zenginleştirme başlıyor — %d chunk | model: %s | concurrency: %d",
        len(docs), LLM_ENRICHMENT_MODEL, _concurrency,
    )

    results: List[LangchainDocument] = list(
        await asyncio.gather(*[_enrich_single(doc, sem, _retries) for doc in docs])
    )

    enriched = sum(1 for r in results if r.metadata.get("llm_enriched"))
    logger.info(
        "Zenginleştirme tamamlandı — ✓ %d başarılı | ✗ %d atlandı/başarısız",
        enriched, len(results) - enriched,
    )
    return results

# ─────────────────────────────────────────────
# İSTATİSTİK
# ─────────────────────────────────────────────
def generate_statistics(documents: List[LangchainDocument]) -> dict:
    """Chunk istatistikleri: kaynak, tip, boyut, mülga, çok-parçalı maddeler."""
    if not documents:
        logger.warning("Belge listesi boş - istatistik oluşturulamadı")
        return {
            "total_chunks":        0,
            "by_source":           {},
            "by_article_type":     {},
            "mulga_chunks":        0,
            "char_count":          {"min": 0, "max": 0, "avg": 0.0},
            "multi_part_articles": [],
            "error":               "No documents provided",
        }

    stats: dict = {
        "total_chunks":        len(documents),
        "by_source":           {},
        "by_article_type":     {},
        "mulga_chunks":        0,
        "char_count":          {"min": float("inf"), "max": 0, "avg": 0.0},
        "multi_part_articles": [],
    }
    total_chars = 0
    seen_multi: Dict[str, int] = {}

    for d in documents:
        m   = d.metadata
        cc  = m.get("char_count", len(d.page_content))
        src = m.get("source", "?")
        atp = m.get("article_type", "?")

        stats["by_source"][src]       = stats["by_source"].get(src, 0) + 1
        stats["by_article_type"][atp] = stats["by_article_type"].get(atp, 0) + 1
        total_chars                  += cc
        stats["char_count"]["min"]    = min(stats["char_count"]["min"], cc)
        stats["char_count"]["max"]    = max(stats["char_count"]["max"], cc)

        if m.get(MetadataKey.IS_MULGA_SOURCE):
            stats["mulga_chunks"] += 1

        ct  = m.get("chunk_total", 1)
        art = m.get("article_reference", "")
        if ct > 1 and art and art not in seen_multi:
            seen_multi[art] = ct

    stats["char_count"]["avg"] = round(total_chars / len(documents), 1)
    if stats["char_count"]["min"] == float("inf"):
        stats["char_count"]["min"] = 0

    stats["multi_part_articles"] = [
        {"article": k, "parts": v}
        for k, v in sorted(seen_multi.items(), key=lambda x: x[1], reverse=True)
    ]

    logger.info(
        "İstatistik: %d chunk, %d kaynak, %d mülga, %d çok parçalı madde",
        stats["total_chunks"], len(stats["by_source"]),
        stats["mulga_chunks"], len(stats["multi_part_articles"]),
    )
    return stats

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    from src.config import STATS_FILE_PATH

    os.makedirs(str(PROCESSED_DIR), exist_ok=True)
    logger.info("### TÜRK İŞ HUKUKU RAG — CHUNKLAMA ###")

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
        datefmt="%H:%M:%S",
    )
    
    resolver = LawMappingResolver(MAPPING_FILE_PATH)
    docs     = load_all_documents(str(RAW_DATA_DIR), resolver)

    if not docs:
        logger.error("Belge işlenemedi.")
        sys.exit(1)

    docs  = asyncio.run(enrich_chunks(docs))
    stats = generate_statistics(docs)

    logger.info("Toplam chunk      : %d", stats["total_chunks"])
    logger.info("Mülga chunk       : %d", stats["mulga_chunks"])
    logger.info(
        "Char min/max/avg  : %d / %d / %.1f",
        stats["char_count"]["min"],
        stats["char_count"]["max"],
        stats["char_count"]["avg"],
    )
    logger.info("Kaynak dağılımı:")
    for src, cnt in sorted(stats["by_source"].items()):
        logger.info("  %-50s %d", src, cnt)

    multi = stats["multi_part_articles"]
    if multi:
        logger.info(
            "Çok parçalı (%d madde) — en parçalanmış %d:",
            len(multi), max((x["parts"] for x in multi), default=0),
        )
        for x in multi[:5]:
            logger.info("  - %s: %d parça", x["article"], x["parts"])

    with open(JSON_FILE_PATH, "w", encoding="utf-8") as f:
        json.dump(
            [{"page_content": d.page_content, "metadata": d.metadata} for d in docs],
            f, ensure_ascii=False, indent=2,
        )
    with open(STATS_FILE_PATH, "w", encoding="utf-8") as f:
        json.dump(stats, f, ensure_ascii=False, indent=2)

    logger.info("=" * 70)
    logger.info("✓ Chunklama tamamlandı: %s (%d chunk)", JSON_FILE_PATH, len(docs))
    logger.info("  İstatistik: %s", STATS_FILE_PATH)
    logger.info("  Sonraki adım: python -m src.create_vector_db")
    logger.info("=" * 70)